from flask import Flask, request, jsonify, render_template, redirect, url_for, make_response, send_file, flash
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
import os
import time
import io
from docx import Document
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import csv
from reportlab.pdfgen import canvas


app = Flask(__name__)

#Configure Mariadb Connection
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://thato:90059Jay#@127.0.0.1:3306/DoForms'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'your-secret-key-here'  # Change this to a random secret key

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

db = SQLAlchemy(app)

# Configure upload folder
UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit

# Create database models here:
# Update User model to include password and inherit from UserMixin
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(1000))
    profile_picture = db.Column(db.String(120))
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

# Create technical installation table here:
class Technical_Installation(db.Model):
    TechnicalInstallationID = db.Column(db.Integer, primary_key=True)
    CustomerName = db.Column(db.String(120), nullable=False)
    CustomerSite = db.Column(db.String(120), nullable=False)
    Date = db.Column(db.String(120), default=lambda: datetime.now().strftime("%d/%m/%Y"))
    Time = db.Column(db.String(120), default=lambda: datetime.now().strftime("%I:%M%p"))
    JobNumber = db.Column(db.String(120))
    InvoiceNumber = db.Column(db.String(120))
    plantMake = db.Column(db.String(120))
    plantModel = db.Column(db.String(120))
    SerialNumber = db.Column(db.String(120))
    EngineMake = db.Column(db.String(120))
    EngineModel = db.Column(db.String(120))
    EngineSerialNumber = db.Column(db.String(120))
    PlantNumber = db.Column(db.String(120))
    PlantKmsHrs = db.Column(db.String(120))
    WarrantyStartDate = db.Column(db.String(120))
    WarrantyEndDate = db.Column(db.String(120))
    InstallationPhoto1 = db.Column(db.String(120))
    InstallationPhoto2 = db.Column(db.String(120))
    InstallationPhoto3 = db.Column(db.String(120))
    InstallationPhoto4 = db.Column(db.String(120))
    InstallationPhoto5 = db.Column(db.String(120))
    InstallationPhoto6 = db.Column(db.String(120))
    InstallationPhoto7 = db.Column(db.String(120))
    InstallationPhoto8 = db.Column(db.String(120))
    InstallationPhoto9 = db.Column(db.String(120))
    InstallationPhoto10 = db.Column(db.String(120))
    AdditionalCommentsSuggestion = db.Column(db.String(120))
    CustomerFullName = db.Column(db.String(120), nullable=False)
    CustomerDate = db.Column(db.String(120))
    TechnitianFullName = db.Column(db.String(120), nullable=False)
    TechnitianDate = db.Column(db.String(120))
    EmailRepeortTo = db.Column(db.String(120))

    def __repr__(self) -> str:
        return f"{self.TechnicalInstallationID} {self.CustomerName}"

# Create Exhaust Inspection Table Here:
class Exhaust_Inspection(db.Model):
    ExhaustInspectionID = db.Column(db.Integer, primary_key=True)
    CustomerName = db.Column(db.String(120), nullable=False)
    CustomerSite = db.Column(db.String(120))
    Date = db.Column(db.String(120), default=lambda: datetime.now().strftime("%d/%m/%Y"))
    Time = db.Column(db.String(120), default=lambda: datetime.now().strftime("%I:%M:%p"))
    ContactPerson = db.Column(db.String(120), nullable=False)
    InspectedBy = db.Column(db.String(120), nullable=False)
    InspectionAssistant = db.Column(db.String(120))
    PlantNumber = db.Column(db.String(120))
    PlantMake = db.Column(db.String(120))
    plantKmsHrs = db.Column(db.String(120))
    PlantModel = db.Column(db.String(120))
    SerialNumber = db.Column(db.String(120))
    EngineMake = db.Column(db.String(120))
    EngineModel = db.Column(db.String(120))
    EngineSerialNumber = db.Column(db.String(120))
    ExhaustMufflerPartNumber = db.Column(db.String(120))
    PhotoOfExhaustMuffler = db.Column(db.String(120))
    CommentsOnExhaustMuffler = db.Column(db.String(120))
    ExhaustPipePartNumber1 = db.Column(db.String(120))
    PhotoOfExhaustPipe1 = db.Column(db.String(120))
    CommentsOnExhaustPipe1 = db.Column(db.String(120))
    ExhaustPipePartNumber2 = db.Column(db.String(120))
    PhotoOfExhaustPipe2 = db.Column(db.String(120))
    CommentsOnExhaustPipe2 = db.Column(db.String(120))
    ExhaustPipePartNumber3 = db.Column(db.String(120))
    PhotoOfExhaustPipe3 = db.Column(db.String(120))
    CommentsOnExhaustPipe3 = db.Column(db.String(120))
    ExhaustPipePartNumber4 = db.Column(db.String(120))
    PhotoOfExhaustPipe4 = db.Column(db.String(120))
    CommentsOnExhaustPipe4 = db.Column(db.String(120))
    ExhaustPipePartNumber5 = db.Column(db.String(120))
    PhotoOfExhaustPipe5 = db.Column(db.String(120))
    CommentsOnExhaustPipe5 = db.Column(db.String(120))
    ExhaustPipePartNumber6 = db.Column(db.String(120))
    PhotoOfExhaustPipe6 = db.Column(db.String(120))
    CommentsOnExhaustPipe6 = db.Column(db.String(120))
    ExhaustPipePartNumber7 = db.Column(db.String(120))
    PhotoOfExhaustPipe7 = db.Column(db.String(120))
    CommentsOnExhaustPipe7 = db.Column(db.String(120))
    ExhaustPipePartNumber8 = db.Column(db.String(120))
    PhotoOfExhaustPipe8 = db.Column(db.String(120))
    CommentsOnExhaustPipe8 = db.Column(db.String(120))
    GeneralComments = db.Column(db.String(120))
    EmailReportTo = db.Column(db.String(120))

    def __repr__(self) -> str:
        return f"{self.ExhaustInspectionID} {self.CustomerName}"

class Plant_Inspection(db.Model):
    PlantInspectionID = db.Column(db.Integer, primary_key=True)
    CustomerName = db.Column(db.String(120), nullable=False)
    CustomerSite = db.Column(db.String(120), nullable=False)
    Date = db.Column(db.String(120), default=lambda: datetime.now().strftime("%d/%m/%Y"))
    Time = db.Column(db.String(120), default=lambda: datetime.now().strftime("%I:%M:%p"))
    ContactPerson = db.Column(db.String(120))
    InspectedBy = db.Column(db.String(120))
    InspectionAssistant = db.Column(db.String(120))
    PhotoOfEquipmentInspected = db.Column(db.LargeBinary)
    PlantNumber = db.Column(db.String(120))
    PlantMake = db.Column(db.String(120))
    PlantModel = db.Column(db.String(120))
    PlantKmsHrs = db.Column(db.String(120))
    SerialNumber = db.Column(db.String(120))
    EngineMake = db.Column(db.String(120))
    EngineModel = db.Column(db.String(120))
    EngineSerialNumber = db.Column(db.String(120))
    GeneralComments = db.Column(db.Text)
    EmailReportTo = db.Column(db.String(120))
    
    def __repr__(self) -> str:
        return f"{self.PlantInspectionID} {self.CustomerName}"
        
class Fuel_System(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    PrimaryFuelFilterPartNumber = db.Column(db.String(80))
    PrimaryFuelFilterPhoto = db.Column(db.LargeBinary)
    CommentsOnFuelFilter = db.Column(db.Text)
    SecondaryFuelFilterPartNumber = db.Column(db.String(80))
    SecondaryFuelFilterPhoto = db.Column(db.LargeBinary)
    CommentsOnSecondaryFuelFilter = db.Column(db.Text)
    FuelWaterSeparatorPartNumber = db.Column(db.String(80))
    FuelWaterSeparatorPhoto = db.Column(db.LargeBinary)
    CommentsOnFuelWaterSeparator = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"
        
class Lube_System(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    OilFilterPartNumber = db.Column(db.String(80))
    OilFilterPhoto = db.Column(db.LargeBinary)
    CommentsOnOilFilter = db.Column(db.Text)
    OilFilterPartNumber1 = db.Column(db.String(80))
    OilFilterPhoto1 = db.Column(db.LargeBinary)
    CommentsOnOilFilter1 = db.Column(db.Text)
    OilFilterPartNumber2 = db.Column(db.String(80))
    OilFilterPhoto2 = db.Column(db.LargeBinary)
    CommentsOnOilFilter2 = db.Column(db.Text)
    OilFilterPartNumber3 = db.Column(db.String(80))
    OilFilterPhoto3 = db.Column(db.LargeBinary)
    CommentsOnOilFilter3 = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class Bypass_Oil_Filter(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    BypassOilFilterPartNumber = db.Column(db.String(80))
    BypassOilFilterPhoto = db.Column(db.LargeBinary)
    CommentsOnBypassOilFilter = db.Column(db.Text)
    BypassOilFilterPartNumber1 = db.Column(db.String(80))
    BypassOilFilterPhoto1 = db.Column(db.LargeBinary)
    CommentsOnBypassOilFilter1 = db.Column(db.Text)
    BypassOilFilterPartNumber2 = db.Column(db.String(80))
    BypassOilFilterPhoto2 = db.Column(db.LargeBinary)
    CommentsOnBypassOilFilter2 = db.Column(db.Text)
    BypassOilFilterPartNumber3 = db.Column(db.String(80))
    BypassOilFilterPhoto3 = db.Column(db.LargeBinary)
    CommentsOnBypassOilFilter3 = db.Column(db.Text)
    
    def __repr__(self) -> str:
        return f"{self.InspectionID}"
        
class Hydraulic_System(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    HydraulicFilterPartNumber = db.Column(db.String(80))
    HydraulicFilterPhoto = db.Column(db.LargeBinary)
    CommentsOnHydraulicFilter = db.Column(db.Text)
    HydraulicFilterPartNumber1 = db.Column(db.String(80))
    HydraulicFilterPhoto1 = db.Column(db.LargeBinary)
    CommentsOnHydraulicFilter1 = db.Column(db.Text)
    HydraulicFilterPartNumber2 = db.Column(db.String(80))
    HydraulicFilterPhoto2 = db.Column(db.LargeBinary)
    CommentsOnHydraulicFilter2 = db.Column(db.Text)
    HydraulicFilterPartNumber3 = db.Column(db.String(80))
    HydraulicFilterPhoto3 = db.Column(db.LargeBinary)
    CommentsOnHydraulicFilter3 = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"
        
class PowerTrain(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    TransmissionFilterPartNumber = db.Column(db.String(80))
    TransmissionFilterPhoto = db.Column(db.LargeBinary)
    CommentsOnTransmissionFilter = db.Column(db.Text)
    TransmissionFilterPartNumber1 = db.Column(db.String(80))
    TransmissionFilterPhoto1 = db.Column(db.LargeBinary)
    CommentsOnTransmissionFilter1 = db.Column(db.Text)
    TransmissionFilterPartNumber2 = db.Column(db.String(80))
    TransmissionFilterPhoto2 = db.Column(db.LargeBinary)
    CommentsOnTransmissionFilter2 = db.Column(db.Text)
    TransmissionFilterPartNumber3 = db.Column(db.String(80))
    TransmissionFilterPhoto3 = db.Column(db.LargeBinary)
    CommentsOnTransmissionFilter3 = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class CoolingSystem(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    CoolantFilterPartNumber = db.Column(db.String(80))
    CoolantFilterPhoto = db.Column(db.LargeBinary)
    CommentsOnCoolantFilter = db.Column(db.Text)
    CoolantFilterPartNumber1 = db.Column(db.String(80))
    CoolantFilterPhoto1 = db.Column(db.LargeBinary)
    CommentsOnCoolantFilter1 = db.Column(db.Text)    
    CoolantFilterPartNumber2 = db.Column(db.String(80))
    CoolantFilterPhoto2 = db.Column(db.LargeBinary)
    CommentsOnCoolantFilter2 = db.Column(db.Text)
    CoolantFilterPartNumber3 = db.Column(db.String(80))
    CoolantFilterPhoto3 = db.Column(db.LargeBinary)
    CommentsOnCoolantFilter3 = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"
        
class BreatherFilter(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    BreatherFilterPartNumber = db.Column(db.String(80))
    BreatherFilterPhoto = db.Column(db.LargeBinary)
    CommentsOnBreatherFilter = db.Column(db.Text)
    BreatherFilterPartNumber1 = db.Column(db.String(80))
    BreatherFilterPhoto1 = db.Column(db.LargeBinary)
    CommentsOnBreatherFilter1 = db.Column(db.Text)
    BreatherFilterPartNumber2 = db.Column(db.String(80))
    BreatherFilterPhoto2 = db.Column(db.LargeBinary)
    CommentsOnBreatherFilter2 = db.Column(db.Text)
    BreatherFilterPartNumber3 = db.Column(db.String(80))
    BreatherFilterPhoto3 = db.Column(db.LargeBinary)
    CommentsOnBreatherFilter3 = db.Column(db.Text)
    
    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class AirDryer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    AirDryerPartNumber = db.Column(db.String(80))
    AirDryerPhoto = db.Column(db.LargeBinary)
    CommentsOnAirDryer = db.Column(db.Text)
    AirDryerPartNumber1 = db.Column(db.String(80))
    AirDryerPhoto1 = db.Column(db.LargeBinary)
    CommentsOnAirDryer1 = db.Column(db.Text)
    AirDryerPartNumber2 = db.Column(db.String(80))
    AirDryerPhoto2 = db.Column(db.LargeBinary)
    CommentsOnAirDryer2 = db.Column(db.Text)
    AirDryerPartNumber3 = db.Column(db.String(80))
    AirDryerPhoto3 = db.Column(db.LargeBinary)
    CommentsOnAirDryer3 = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class CabinFilter(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    CabinFilterPartNumber = db.Column(db.String(80))
    CabinFilterPhoto = db.Column(db.LargeBinary)
    CommentsOnCabinFilter = db.Column(db.Text)
    CabinFilterPartNumber1 = db.Column(db.String(80))
    CabinFilterPhoto1 = db.Column(db.LargeBinary)
    CommentsOnCabinFilter1 = db.Column(db.Text)
    CabinFilterPartNumber2 = db.Column(db.String(80))
    CabinFilterPhoto2 = db.Column(db.LargeBinary)
    CommentsOnCabinFilter2 = db.Column(db.Text)
    CabinFilterPartNumber3 = db.Column(db.String(80))
    CabinFilterPhoto3 = db.Column(db.LargeBinary)
    CommentsOnCabinFilter3 = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class AirFiltration(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    AirCleanerModel = db.Column(db.String(80))
    NumberOfAirCleanerUnits = db.Column(db.String(80))
    AirCleanerArrangementPhoto = db.Column(db.LargeBinary)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class Air_Cleaner_Element(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    PrimaryElementPartNumber = db.Column(db.String(80))
    PrimaryElementDateCode = db.Column(db.String(80))
    PrimaryAirCleanerElementPhoto = db.Column(db.LargeBinary)
    CommentsOnPrimaryElement = db.Column(db.Text)
    
    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class Safety_Air_Cleaner_Element(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))
    SafetyElementPartNumber = db.Column(db.String(80))
    SafetyElementDateCode = db.Column(db.String(80))
    CommentsOnSafetyElement = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class Air_Cleaner_Outlet(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))    
    AirCleanerOutletPartNumber = db.Column(db.String(80))
    AirCleanerOutletDateCode = db.Column(db.String(80))
    AirCleanerOutletPhoto = db.Column(db.LargeBinary)
    CommentsOnAirCleanerOutlet = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class Clean_Air_Ducting(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))    
    CleanAirDuctingPhoto = db.Column(db.LargeBinary)
    CommentsOnCleanAirDucting = db.Column(db.Text)
    
    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class Restriction_Indicator(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))    
    RestrictionIndicatorPartNumber = db.Column(db.String(80))
    RestrictionIndicatorValue = db.Column(db.String(80))
    RestrictionIndicatorPhoto = db.Column(db.LargeBinary)
    CommentsOnRestrictionIndicator = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class Vacuator_Valves(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))    
    VacuatorValvesPartNumber = db.Column(db.String(80))
    VacuatorValvesPhoto = db.Column(db.LargeBinary)
    CommentsOnVacuatorValves = db.Column(db.Text)    

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

class Other_Observation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    InspectionID = db.Column(db.Integer, db.ForeignKey('plant__inspection.PlantInspectionID'))    
    OtherObservationPhoto = db.Column(db.LargeBinary)
    CommentsOnOtherObservation = db.Column(db.Text)
    OtherObservationPhoto1 = db.Column(db.LargeBinary)
    CommentsOnOtherObservation1 = db.Column(db.Text)
    OtherObservationPhoto2 = db.Column(db.LargeBinary)
    CommentsOnOtherObservation2 = db.Column(db.Text)
    OtherObservationPhoto3 = db.Column(db.LargeBinary)
    CommentsOnOtherObservation3 = db.Column(db.Text)
    OtherObservationPhoto4 = db.Column(db.LargeBinary)
    CommentsOnOtherObservation4 = db.Column(db.Text)

    def __repr__(self) -> str:
        return f"{self.InspectionID}"

# User loader callback
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Helper function to check allowed file extensions
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Create tables (run this once)
with app.app_context():
    db.drop_all()  # This will delete all tables
    db.create_all()  # This will recreate them with the current schema

# Helper function to handle file upload    
def save_profile_picture(file):
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        # Add timestamp to make filename unique
        unique_filename = f"{int(time.time())}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        return unique_filename
    return None

# Add new routes for authentication
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('get_all_users_page'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            login_user(user)
            return redirect(url_for('get_all_users_page'))
        flash('Invalid username or password')
    return render_template('login.html')

@app.route('/logout')
@login_required  # This ensures only logged-in users can access this route
def logout():
    logout_user()  # This clears the user session
    flash('You have been logged out successfully.', 'success')  # Optional: Add a flash message
    return redirect(url_for('login'))  # Redirect to login page after logout
    
@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('get_all_users_page'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        
        if User.query.filter_by(username=username).first():
            flash('Username already exists')
            return redirect(url_for('register'))
        
        if User.query.filter_by(email=email).first():
            flash('Email already exists')
            return redirect(url_for('register'))
        
        new_user = User(username=username, email=email)
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()
        
        flash('Registration successful! Please login.')
        return redirect(url_for('login'))
    
    return render_template('register.html')

# Protect existing routes with @login_required
# API Routes
# Place export logic/functionality here:
@app.route('/export/<format>', methods=['POST'])
def export_surveys(format):
    selected_ids = request.form.get('selected_ids', '').split(',')
    
    # Get the selected surveys from your database
    surveys = []  # Replace with your actual query
    # Example: surveys = Survey.query.filter(Survey.id.in_(selected_ids)).all()
    
    if format == 'excel':
        # Create CSV (Excel-compatible)
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Write header
        writer.writerow(['ID', 'Type', 'Date', 'Inspector', 'Status'])
        
        # Write data
        for survey in surveys:
            writer.writerow([
                survey.id,
                survey.type,
                survey.date.strftime('%Y-%m-%d'),
                survey.inspector,
                survey.status
            ])
        
        response = make_response(output.getvalue())
        response.headers['Content-Disposition'] = 'attachment; filename=surveys_export.csv'
        response.headers['Content-type'] = 'text/csv'
        return response
    
    elif format == 'word':
        # Create Word document
        document = Document()
        document.add_heading('Surveys Export', 0)
        
        # Add table
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Date'
        hdr_cells[3].text = 'Inspector'
        hdr_cells[4].text = 'Status'
        
        # Data rows
        for survey in surveys:
            row_cells = table.add_row().cells
            row_cells[0].text = str(survey.id)
            row_cells[1].text = survey.type
            row_cells[2].text = survey.date.strftime('%Y-%m-%d')
            row_cells[3].text = survey.inspector
            row_cells[4].text = survey.status
        
        # Save to buffer
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        
        response = make_response(output.read())
        response.headers['Content-Disposition'] = 'attachment; filename=surveys_export.docx'
        response.headers['Content-type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response
    
    elif format == 'pdf':
        # Create PDF
        output = io.BytesIO()
        p = canvas.Canvas(output)
        
        # Set title
        p.setFont("Helvetica-Bold", 16)
        p.drawString(100, 800, "Surveys Export")
        
        # Set table headers
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, 770, "ID")
        p.drawString(100, 770, "Type")
        p.drawString(200, 770, "Date")
        p.drawString(300, 770, "Inspector")
        p.drawString(400, 770, "Status")
        
        # Draw data
        p.setFont("Helvetica", 10)
        y = 750
        for survey in surveys:
            p.drawString(50, y, str(survey.id))
            p.drawString(100, y, survey.type)
            p.drawString(200, y, survey.date.strftime('%Y-%m-%d'))
            p.drawString(300, y, survey.inspector)
            p.drawString(400, y, survey.status)
            y -= 20
            if y < 50:  # Add new page if we're at the bottom
                p.showPage()
                y = 800
        
        p.save()
        output.seek(0)
        
        response = make_response(output.read())
        response.headers['Content-Disposition'] = 'attachment; filename=surveys_export.pdf'
        response.headers['Content-type'] = 'application/pdf'
        return response
    
    return "Invalid export format", 400


@app.route('/users', methods=['POST'])
@login_required
def create_user():
    try:
        username = request.form.get('username')
        email = request.form.get('email')
        profile_pic = request.files.get('profile_picture')
        
        if not username or not email:
            return jsonify({'error': 'Username and email are required'}), 400
        
        # Check for existing user
        if User.query.filter_by(username=username).first():
            return jsonify({'error': 'Username already exists'}), 409
        if User.query.filter_by(email=email).first():
            return jsonify({'error': 'Email already exists'}), 409
        
        # Handle profile picture upload
        profile_pic_filename = None
        if profile_pic:
            profile_pic_filename = save_profile_picture(profile_pic)
            if not profile_pic_filename:
                return jsonify({'error': 'Invalid file type'}), 400
        
        # Create new user
        new_user = User(
            username=username,
            email=email,
            profile_picture=profile_pic_filename
        )
        
        db.session.add(new_user)
        db.session.commit()
        
        return jsonify({
            'message': 'User created successfully',
            'user': {
                'id': new_user.id,
                'username': new_user.username,
                'email': new_user.email,
                'profile_picture': new_user.profile_picture
            }
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500
                
@app.route('/users', methods=['GET'])
@login_required
def get_all_users_api():
    users = User.query.all()
    output = []
    for user in users:
        user_data = {'id': user.id, 'username': user.username, 'email': user.email}
        output.append(user_data)
    return jsonify({'users': output})

@app.route('/users/<int:user_id>', methods=['GET'])
@login_required
def get_user(user_id):
    user = User.query.get_or_404(user_id)
    return jsonify({'id': user.id, 'username': user.username, 'email': user.email})

@app.route('/users/<int:user_id>', methods=['PUT'])
@login_required
def update_user(user_id):
    user = User.query.get_or_404(user_id)
    
    try:
        username = request.form.get('username')
        email = request.form.get('email')
        profile_pic = request.files.get('profile_picture')
        
        # Update username if changed
        if username and username != user.username:
            if User.query.filter_by(username=username).first():
                return jsonify({'error': 'Username already taken'}), 409
            user.username = username
        
        # Update email if changed
        if email and email != user.email:
            if User.query.filter_by(email=email).first():
                return jsonify({'error': 'Email already taken'}), 409
            user.email = email
        
        # Handle profile picture update
        if profile_pic:
            # Delete old picture if exists
            if user.profile_picture:
                old_pic_path = os.path.join(app.config['UPLOAD_FOLDER'], user.profile_picture)
                if os.path.exists(old_pic_path):
                    os.remove(old_pic_path)
            
            # Save new picture
            profile_pic_filename = save_profile_picture(profile_pic)
            if profile_pic_filename:
                user.profile_picture = profile_pic_filename

        # Inside the update_user route
        if request.form.get('remove_picture') == 'on' and user.profile_picture:
            old_pic_path = os.path.join(app.config['UPLOAD_FOLDER'], user.profile_picture)
            if os.path.exists(old_pic_path):
                os.remove(old_pic_path)
            user.profile_picture = None
        
        db.session.commit()
        
        return jsonify({
            'message': 'User updated successfully',
            'user': {
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'profile_picture': user.profile_picture
            }
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500
        
        
@app.route('/users/<int:user_id>', methods=['DELETE'])
@login_required
def delete_user(user_id):
    user = User.query.get_or_404(user_id)
    db.session.delete(user)
    db.session.commit()
    return jsonify({'message': 'User deleted successfully'})

# HTML Routes

@app.route('/')
@login_required

#def get_all_users_page():
#    users = User.query.all()
#    return render_template('users.html', users=users)

@app.route('/')
def dashboard():
    # Pagination parameters
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 5, type=int)  # Default to 5 items per page
    
    # Get total counts (keep your existing code)
    total_surveys = 124
    plant_inspections = 42
    exhaust_inspections = 35
    technical_installations = 28
    birth_certificates = 19
    
    # Get paginated surveys (replace with your actual query)
    # Example with SQLAlchemy:
    # surveys = Survey.query.order_by(Survey.date.desc()).paginate(page=page, per_page=per_page)
    
    # For demo purposes, we'll create mock paginated data
    all_surveys = [
        {
            'id': i,
            'type': random.choice(['Plant Inspection', 'Exhaust Inspection', 'Technical Installation', 'Birth Certificate']),
            'date': datetime.now() - timedelta(days=random.randint(0, 30)),
            'inspector': f"Inspector {random.randint(1, 5)}",
            'status': random.choice(['Completed', 'In Progress'])
        }
        for i in range(1, 21)  # 20 mock surveys
    ]
    
    # Manual pagination for demo
    start = (page - 1) * per_page
    end = start + per_page
    paginated_surveys = all_surveys[start:end]
    
    # Generate monthly data for the chart
    monthly_labels = [datetime.now().replace(day=1, month=i).strftime('%b') for i in range(1, 13)]
    monthly_data = [random.randint(5, 20) for _ in range(12)]
    
    return render_template(
        'dashboard.html',
        total_surveys=total_surveys,
        plant_inspections=plant_inspections,
        exhaust_inspections=exhaust_inspections,
        technical_installations=technical_installations,
        birth_certificates=birth_certificates,
        recent_surveys=paginated_surveys,
        monthly_labels=monthly_labels,
        monthly_data=monthly_data,
        page=page,
        per_page=per_page,
        total_pages=(len(all_surveys) + per_page - 1) // per_page  # Calculate total pages
    )


@app.route('/about')
@login_required
def about_us():
    return render_template('about.html')

@app.route('/privacy')
@login_required
def privacy_policy():
    return render_template('privacy.html')

@app.route('/plant')
@login_required
def plant_inspection():
    return render_template('plant.html')

@app.route('/exhaust')
@login_required
def exhaust_inspection():
    return render_template('exhaust.html')

@app.route('/technical')
@login_required
def technical_installation():
    return render_template('technical.html')

@app.route('/certificates')
@login_required
def birth_certificate():
    return render_template('certificate.html')

@app.route('/users/add')
@login_required
def add_user_page():
    return render_template('add_user.html')

@app.route('/users/<int:user_id>/edit')
@login_required
def edit_user_page(user_id):
    user = User.query.get_or_404(user_id)
    return render_template('edit_user.html', user=user)

@app.route('/export/word')
@login_required
def export_word():
    users = User.query.all()
    
    document = Document()
    document.add_heading('User Management System Report', 0)
    
    # Add a table
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Username'
    hdr_cells[2].text = 'Email'
    hdr_cells[3].text = 'Profile Picture'
    
    # Add data rows
    for user in users:
        row_cells = table.add_row().cells
        row_cells[0].text = str(user.id)
        row_cells[1].text = user.username
        row_cells[2].text = user.email
        row_cells[3].text = user.profile_picture if user.profile_picture else 'None'
    
    # Save to a BytesIO buffer
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name='users_report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/export/excel')
@login_required
def export_excel():
    users = User.query.all()
    
    # Create a DataFrame
    data = {
        'ID': [user.id for user in users],
        'Username': [user.username for user in users],
        'Email': [user.email for user in users],
        'Profile Picture': [user.profile_picture if user.profile_picture else 'None' for user in users]
    }
    df = pd.DataFrame(data)
    
    # Create output buffer
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    df.to_excel(writer, sheet_name='Users', index=False)
    writer.close()
    output.seek(0)
    
    return send_file(
        output,
        as_attachment=True,
        download_name='users_report.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

@app.route('/export/pdf')
@login_required
def export_pdf():
    users = User.query.all()
    
    # Create a buffer for the PDF
    buffer = io.BytesIO()
    
    # Create the PDF object
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    
    # Add title
    styles = getSampleStyleSheet()
    elements.append(Paragraph("User Management System Report", styles['Title']))
    
    # Prepare data for table
    data = [['ID', 'Username', 'Email', 'Profile Picture']]
    for user in users:
        data.append([
            str(user.id),
            user.username,
            user.email,
            user.profile_picture if user.profile_picture else 'None'
        ])
    
    # Create table
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    elements.append(table)
    doc.build(elements)
    
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name='users_report.pdf',
        mimetype='application/pdf'
    )

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')
